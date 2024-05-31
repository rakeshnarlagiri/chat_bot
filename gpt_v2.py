import os
import time

import streamlit as st
from langchain_text_splitters import CharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.chat_models import ChatOpenAI
from langchain.memory import ConversationBufferMemory, StreamlitChatMessageHistory
from langchain.chains import ConversationalRetrievalChain, LLMChain
from langchain.vectorstores import Chroma
from langchain_core.prompts import PromptTemplate
from langchain_community.document_loaders import PyPDFLoader
import chromadb
from utils import get_available_openai_models
import tempfile
from streaming import StreamHandler
import pandas as pd
import chardet
from io import StringIO
from docx import Document
import json

os.environ["OPENAI_API_KEY"] = "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"
DB_DIRECTORY = "./chroma"

TEMPLATE = (
    "Question: {question}\n\n"
    "Use the following pieces of context to answer the question.\n"
    "If you don't know the answer, just say that you don't know, don't try to make up an answer.\n"
    "----------------\n"
    "{context}"
)


@st.cache_data
def chat_model_list():
    return get_available_openai_models(put_first='gpt-3.5-turbo', filter_by='gpt')


@st.cache_data
def embedding_model_list():
    return get_available_openai_models(filter_by='embedding')


def process_uploaded_file(uploaded_file):
    try:
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getvalue())
        return file_path
    except Exception as e:
        st.error(f"Error processing uploaded file: {e}")


def get_saved_databases():
    dbs = []
    for item in os.listdir(DB_DIRECTORY):
        item_path = os.path.join(DB_DIRECTORY, item)
        if os.path.isdir(item_path):
            dbs.append(item)
    return dbs


class StreamlitChatView:
    def __init__(self, default_embeddings_model="text-embedding-ada-002"):
        st.set_page_config(page_title="RAG ChatGPT", page_icon="📚", layout="wide")
        with st.sidebar:
            st.title("RAG ChatGPT")
            with st.expander("Model parameters"):
                self.model_name = st.selectbox("Model:", options=chat_model_list())
                self.temperature = st.slider("Temperature", min_value=0., max_value=2., value=0.7, step=0.01)
                self.top_p = st.slider("Top p", min_value=0., max_value=1., value=1., step=0.01)
                self.frequency_penalty = st.slider("Frequency penalty", min_value=0., max_value=2., value=0., step=0.01)
                self.presence_penalty = st.slider("Presence penalty", min_value=0., max_value=2., value=0., step=0.01)
            with st.expander("Embeddings parameters"):
                self.embeddings_model_name = st.selectbox("Embeddings model:", options=embedding_model_list(),
                                                          index=embedding_model_list().index(default_embeddings_model))
            with st.expander("Prompts"):
                self.context_prompt = st.text_area("Context prompt", value=TEMPLATE)
            self.input_file = st.file_uploader("Upload File", type=["pdf", "docx", "doc", "json", ".xlsx"])
            self.selected_db = st.sidebar.selectbox("Select saved database:", [""] + get_saved_databases(), index=0)
            self.csv_file = st.file_uploader("Upload csv for questions", type="csv")

        self.user_query = st.chat_input(placeholder="Ask me anything!")

    def add_message(self, message, author):
        with st.chat_message(author):
            st.markdown(message)

    def add_message_stream(self, author: str):
        assert author in ["user", "assistant"]
        return StreamHandler(st.chat_message(author).empty())


def read_csv(view):
    if view.csv_file:
        file_content = view.csv_file.read()

        # Detect the encoding of the file content
        result = chardet.detect(file_content)
        encoding = result['encoding']

        # Convert the byte content to a string buffer using the detected encoding
        decoded_content = file_content.decode(encoding)
        file_buffer = StringIO(decoded_content)

        # Read the CSV content from the string buffer
        df = pd.read_csv(file_buffer)

        return df
    # for index, row in df.iterrows():
    #     value = df.iloc[index, 0]
    #     print(value)


def clear_chat(memory):
    memory.chat_memory.messages = []


def prepare_docs(view):
    file_path = process_uploaded_file(view.input_file)
    _, file_extension = os.path.splitext(file_path)
    if file_extension.lower() == '.pdf':
        loader = PyPDFLoader(file_path)
        data = loader.load()
        return data
    elif file_extension.lower() in ['.docx', '.doc']:
        document = Document(file_path)
        data = []
        for paragraph in document.paragraphs:
            data.append(paragraph.text)
        docs = []
        for text in data:
            document = Document()  # Initialize a document object
            document.page_content = text
            document.metadata = {}  # Assign text content to the page_content attribute
            docs.append(document)
        return docs
    elif file_extension.lower() == '.json':

        # Open the JSON file and read its contents
        with open(file_path, 'r') as json_file:
            # Load the JSON data
            data = json.load(json_file)

        docs = []
        for text in data:
            document = Document()  # Initialize a document object
            document.page_content = text
            document.metadata = {}  # Assign text content to the page_content attribute
            docs.append(document)
        return docs

    elif file_extension.lower() == '.xlsx':
        data = pd.read_excel(file_path)
        docs = []
        for text in data:
            document = Document()  # Initialize a document object
            document.page_content = text
            document.metadata = {}  # Assign text content to the page_content attribute
            docs.append(document)
        return docs
    else:
        raise ValueError("Unsupported file type")


def main():
    view = StreamlitChatView()

    # prompt = PromptTemplate.from_template(view.context_prompt)
    prompt = PromptTemplate(
        input_variables=["chat_history", "context", "question"],
        template=view.context_prompt, )

    msgs = StreamlitChatMessageHistory(key="langchain_messages")
    memory = ConversationBufferMemory(memory_key="chat_history", chat_memory=msgs, output_key='answer',
                                      return_messages=True)

    df = read_csv(view)
    if view.input_file:
        data = prepare_docs(view)

    col1, col2 = st.columns([6, 1])

    with col1:
        st.header(f"Chat with Docs using OpenAI")

    with col2:
        if st.button("Clear ↺"):
            clear_chat(memory)

    if view.input_file and not view.selected_db:
        process_pdf(view, memory, prompt, df, data)
    elif view.selected_db and view.selected_db != "" and not view.input_file:
        process_pdf(view, memory, prompt, df)
    elif view.selected_db and view.input_file:
        st.warning("select only one source( Either upload or select db )", icon="⚠️")
    else:
        # st.success("upload or select db to chat with")
        st.info('upload file or select db to chat with', icon="ℹ️")


def process_pdf(view, memory, prompt, df, data=None):
    embeddings = OpenAIEmbeddings(model=view.embeddings_model_name)
    if data:
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        docs = text_splitter.split_documents(data)
        if os.path.exists(os.path.join(DB_DIRECTORY, view.input_file.name)):
            st.warning('already have a db for this', icon="⚠️")
        else:
            vectorstore = Chroma.from_documents(docs, embeddings, persist_directory=os.path.join(DB_DIRECTORY,
                                                                                                 view.input_file.name))
            model(view, memory, prompt, df, vectorstore)
    elif view.selected_db:
        # memory.chat_memory.messages = []
        selected_db_path = os.path.join(DB_DIRECTORY, view.selected_db)
        vectorstore = Chroma(persist_directory=selected_db_path, embedding_function=embeddings)
        model(view, memory, prompt, df, vectorstore)

    else:
        raise ValueError("Either a PDF file or a selected DB must be provided...")


STREAM = False


def model(view, memory, prompt, df, db):
    llm = ChatOpenAI(
        model_name=view.model_name,
        temperature=view.temperature,
        top_p=view.top_p,
        frequency_penalty=view.frequency_penalty,
        presence_penalty=view.presence_penalty
    )

    # conversation_chain = ConversationalRetrievalChain.from_llm(llm=llm,
    #                                                            retriever=db.as_retriever(),
    #                                                            memory=memory,
    #                                                            verbose=True,
    #                                                            prompt=prompt
    #                                                              )
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm,
        retriever=db.as_retriever(search_kwargs={"k": 3}),
        return_source_documents=True,
        verbose=True,
        chain_type="stuff",
        get_chat_history=lambda h: h,
        combine_docs_chain_kwargs={'prompt': prompt},
        memory=memory,
    )

    for message in memory.chat_memory.messages:
        view.add_message(message.content, 'assistant' if message.type == 'ai' else 'user')

    if view.user_query:
        view.add_message(view.user_query, "user")
        if STREAM:
            st_callback = view.add_message_stream("assistant")
            conversation_chain({"question": view.user_query}, callbacks=[st_callback])
        else:
            response = conversation_chain({"question": view.user_query})
            view.add_message(response['answer'], "assistant")

    elif view.csv_file:
        for index, row in df.iterrows():
            value = df.iloc[index, 0]
            view.add_message(value, "user")
            if STREAM:
                st_callback = view.add_message_stream("assistant")
                conversation_chain({"question": value}, callbacks=[st_callback])
            else:
                response = conversation_chain({"question": value})
                view.add_message(response['answer'], "assistant")


if __name__ == "__main__":
    main()
